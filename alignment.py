from docx import Document
from listing import list_number
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH

doc=Document()
table = doc.add_table(rows=0, cols=2)

row=table.add_row().cells
p=row[0].add_paragraph('left justified text')
p.alignment=WD_ALIGN_PARAGRAPH.LEFT
p=row[1].add_paragraph('right justified text')
p.alignment=WD_ALIGN_PARAGRAPH.RIGHT

doc.save('try.docx')